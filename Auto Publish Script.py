#!/usr/bin/env python3
import os, re, sys, json, time
from pathlib import Path

DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY", "")
INPUT_DIR        = "/Users/paulburg/Vibe_coding/PaulBurg.com/articles"
OUTPUT_DIR       = "/Users/paulburg/Vibe_coding/PaulBurg.com/content/posts"
DEFAULT_TAG      = "AI"
EXCERPT_CHARS    = 800

def read_opening(filepath, max_chars=EXCERPT_CHARS):
    ext = Path(filepath).suffix.lower()
    if ext == ".docx":
        try:
            from docx import Document
            doc = Document(filepath)
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            return text[:max_chars]
        except ImportError:
            print("  Run: pip3 install python-docx"); sys.exit(1)
    elif ext == ".doc":
        try:
            import mammoth
            with open(filepath, "rb") as f:
                return mammoth.extract_raw_text(f).value[:max_chars]
        except ImportError:
            print("  Run: pip3 install mammoth"); sys.exit(1)

def read_full(filepath):
    ext = Path(filepath).suffix.lower()
    if ext == ".docx":
        from docx import Document
        doc, paragraphs = Document(filepath), []
        for para in doc.paragraphs:
            text  = para.text.strip()
            style = para.style.name.lower()
            if not text:               paragraphs.append("")
            elif "heading 1" in style: paragraphs.append("# " + text)
            elif "heading 2" in style: paragraphs.append("## " + text)
            elif "heading 3" in style: paragraphs.append("### " + text)
            elif "list" in style:      paragraphs.append("- " + text)
            else:                      paragraphs.append(text)
        out, prev_blank = [], False
        for line in paragraphs:
            if line == "":
                if not prev_blank: out.append("")
                prev_blank = True
            else:
                out.append(line); prev_blank = False
        return "\n\n".join(out)
    elif ext == ".doc":
        import mammoth
        with open(filepath, "rb") as f:
            return mammoth.convert_to_markdown(f).value

def find_pairs(input_dir):
    folder = Path(input_dir)
    if not folder.exists():
        folder.mkdir(parents=True)
        print("Created articles/ folder. Drop _en.doc + _ru.doc files there and run again.")
        sys.exit(0)
    all_files = list(folder.glob("*.doc")) + list(folder.glob("*.docx"))
    en_files  = [f for f in all_files if re.search(r'[_\-]en$', f.stem, re.I)]
    ru_files  = [f for f in all_files if re.search(r'[_\-]ru$', f.stem, re.I)]
    pairs = []
    for en in en_files:
        base = re.sub(r'[_\-]en$', '', en.stem, flags=re.I)
        ru_match = next((r for r in ru_files
                         if re.sub(r'[_\-]ru$', '', r.stem, flags=re.I).lower() == base.lower()), None)
        if ru_match:
            pairs.append({"base": base, "en": str(en), "ru": str(ru_match)})
        else:
            print("  WARNING: No RU pair for " + en.name + " — skipping")
    return pairs

def extract_metadata(opening_en, opening_ru, base_name):
    suggested_slug = re.sub(r'[^a-z0-9]+', '-', base_name.lower()).strip('-')
    prompt = (
        "Extract metadata from these article openings. "
        "Return ONLY a JSON object, no other text, no markdown fences.\n\n"
        "English opening:\n" + opening_en + "\n\n"
        "Russian opening:\n" + opening_ru + "\n\n"
        "Return exactly this JSON structure:\n"
        "{\n"
        '  "slug": "' + suggested_slug + '",\n'
        '  "title_en": "Full article title in English",\n'
        '  "title_ru": "Full article title in Russian",\n'
        '  "excerpt_en": "One sentence SEO description max 155 chars",\n'
        '  "excerpt_ru": "One sentence SEO description in Russian max 155 chars",\n'
        '  "tag": "AI"\n'
        "}\n\n"
        "Tag must be one of: AI | Business | Automation | CRM\n"
        "Return ONLY the JSON."
    )
    try:
        import requests
        r = requests.post(
            "https://api.deepseek.com/v1/chat/completions",
            headers={"Authorization": "Bearer " + DEEPSEEK_API_KEY, "Content-Type": "application/json"},
            json={"model": "deepseek-chat", "messages": [{"role": "user", "content": prompt}],
                  "max_tokens": 400, "temperature": 0.1},
            timeout=30,
        )
        r.raise_for_status()
        raw = r.json()["choices"][0]["message"]["content"].strip()
        raw = re.sub(r'^```json\s*|^```\s*|\s*```$', '', raw)
        return json.loads(raw)
    except Exception as e:
        print("  WARNING: DeepSeek failed (" + str(e) + "), using filename fallback")
        return {
            "slug":       suggested_slug,
            "title_en":   base_name.replace("-", " ").replace("_", " ").title(),
            "title_ru":   base_name.replace("-", " ").replace("_", " ").title(),
            "excerpt_en": "Article on paulburg.com",
            "excerpt_ru": "Article on paulburg.com",
            "tag":        DEFAULT_TAG,
        }

def build_mdx(title, date_str, tag, excerpt, content):
    t = title.replace('"', '\\"')
    e = excerpt.replace('"', '\\"')
    return "\n".join([
        '---',
        'title: "' + t + '"',
        'date: "' + date_str + '"',
        'tag: "' + tag + '"',
        'excerpt: "' + e + '"',
        '---',
        '',
        content,
    ])

def main():
    from datetime import date
    today = date.today().isoformat()
    print("paulburg.com Auto Blog Publisher")
    print("Input:  " + INPUT_DIR)
    print("Output: " + OUTPUT_DIR + "\n")
    pairs = find_pairs(INPUT_DIR)
    if not pairs:
        print("No EN+RU pairs found. Name files: article_en.doc + article_ru.doc")
        sys.exit(0)
    print("Found " + str(len(pairs)) + " pair(s): " + str([p['base'] for p in pairs]) + "\n")
    for i, pair in enumerate(pairs, 1):
        base = pair["base"]
        print("[" + str(i) + "/" + str(len(pairs)) + "] " + base)
        print("  Reading openings...")
        opening_en = read_opening(pair["en"])
        opening_ru = read_opening(pair["ru"])
        print("  Extracting metadata via DeepSeek...")
        meta    = extract_metadata(opening_en, opening_ru, base)
        slug    = meta["slug"]
        out_dir = Path(OUTPUT_DIR) / slug
        if (out_dir / "en.mdx").exists():
            print("  SKIP: " + slug + " already published.")
            continue
        print("  Converting full content...")
        content_en = read_full(pair["en"])
        content_ru = read_full(pair["ru"])
        out_dir.mkdir(parents=True, exist_ok=True)
        (out_dir / "en.mdx").write_text(
            build_mdx(meta["title_en"], today, meta["tag"], meta["excerpt_en"], content_en), encoding="utf-8")
        (out_dir / "ru.mdx").write_text(
            build_mdx(meta["title_ru"], today, meta["tag"], meta["excerpt_ru"], content_ru), encoding="utf-8")
        print("  ✓ content/posts/" + slug + "/")
        print("    EN: " + meta["title_en"])
        print("    RU: " + meta["title_ru"])
        print("    Tag: " + meta["tag"])
        if i < len(pairs):
            time.sleep(0.5)
    print("\n✓ Done. Now run:")
    print("  git add content/posts/")
    print("  git commit -m 'Add blog posts'")
    print("  git push")

if __name__ == "__main__":
    main()
